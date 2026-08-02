"""
Document Parsers for PDF, DOCX, and TXT files.
Extracts raw text along with structural metadata (page numbers, section headers).
Uses a clean registry pattern for easy format extension.
"""

import os
from typing import List, Dict, Any, Callable, Protocol
from dataclasses import dataclass, field
import fitz  # PyMuPDF
import docx


@dataclass
class ExtractedPage:
    """Represents text extracted from a page or section of a document."""
    page_number: int
    text: str
    section_header: str = ""


@dataclass
class ParsedDocument:
    """Represents a parsed document containing pages/sections and metadata."""
    doc_id: str
    filename: str
    file_type: str
    pages: List[ExtractedPage] = field(default_factory=list)
    full_text: str = ""


class DocumentParser(Protocol):
    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        ...


class PDFParser:
    """Extracts text and page numbers from PDF files using PyMuPDF."""
    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        doc = fitz.open(file_path)
        pages: List[ExtractedPage] = []
        full_text_chunks: List[str] = []

        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if text:
                pages.append(ExtractedPage(page_number=page_num, text=text))
                full_text_chunks.append(text)

        doc.close()
        filename = os.path.basename(file_path)
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            file_type="pdf",
            pages=pages,
            full_text="\n\n".join(full_text_chunks),
        )


class DOCXParser:
    """Extracts text and section headings from Microsoft Word DOCX files using python-docx."""
    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        doc = docx.Document(file_path)
        pages: List[ExtractedPage] = []
        full_text_chunks: List[str] = []
        current_header = "Document Start"
        current_page_num = 1
        current_section_text: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if paragraph is a heading
            if para.style and para.style.name.startswith("Heading"):
                if current_section_text:
                    section_str = "\n".join(current_section_text)
                    pages.append(
                        ExtractedPage(
                            page_number=current_page_num,
                            text=section_str,
                            section_header=current_header,
                        )
                    )
                    full_text_chunks.append(section_str)
                    current_page_num += 1
                    current_section_text = []
                current_header = text
            else:
                current_section_text.append(text)

        if current_section_text:
            section_str = "\n".join(current_section_text)
            pages.append(
                ExtractedPage(
                    page_number=current_page_num,
                    text=section_str,
                    section_header=current_header,
                )
            )
            full_text_chunks.append(section_str)

        filename = os.path.basename(file_path)
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            file_type="docx",
            pages=pages,
            full_text="\n\n".join(full_text_chunks),
        )


class TXTParser:
    """Extracts text from plain text files."""
    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"TXT file not found: {file_path}")

        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read().strip()

        filename = os.path.basename(file_path)
        page = ExtractedPage(page_number=1, text=content, section_header="Main Content")
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            file_type="txt",
            pages=[page],
            full_text=content,
        )


class ParserRegistry:
    """Format-to-parser registry for extensible file format handling."""
    def __init__(self):
        self._parsers: Dict[str, DocumentParser] = {
            ".pdf": PDFParser(),
            ".docx": DOCXParser(),
            ".txt": TXTParser(),
        }

    def register_parser(self, extension: str, parser: DocumentParser):
        """Registers a new parser for a specific file extension."""
        ext = extension.lower() if extension.startswith(".") else f".{extension.lower()}"
        self._parsers[ext] = parser

    def get_parser(self, file_path: str) -> DocumentParser:
        """Retrieves appropriate parser based on file extension."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self._parsers:
            supported = ", ".join(self._parsers.keys())
            raise ValueError(
                f"Unsupported file format '{ext}'. Supported formats: {supported}"
            )
        return self._parsers[ext]

    def parse_file(self, file_path: str, doc_id: str) -> ParsedDocument:
        """Parses a file using the registered parser for its extension."""
        parser = self.get_parser(file_path)
        return parser.parse(file_path, doc_id)


# Global default registry instance
global_parser_registry = ParserRegistry()
