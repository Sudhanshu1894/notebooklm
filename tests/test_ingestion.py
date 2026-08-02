"""
Unit Test Suite for GraphRAG Ingestion Pipeline (Parsers, Chunker, Registry).
"""

import os
import pytest
import fitz  # PyMuPDF
import docx
from ingestion.parsers import ParserRegistry, TXTParser, PDFParser, DOCXParser
from ingestion.chunker import chunk_parsed_document
from ingestion.registry import DocumentRegistry


@pytest.fixture
def sample_txt_file(tmp_path):
    txt_path = os.path.join(tmp_path, "sample_doc.txt")
    text_content = (
        "GraphRAG is a hybrid retrieval architecture. "
        "It combines vector search and knowledge graph traversal. "
        "This approach helps answer complex multi-hop queries accurately."
    )
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(text_content)
    return txt_path


@pytest.fixture
def sample_docx_file(tmp_path):
    docx_path = os.path.join(tmp_path, "sample_doc.docx")
    doc = docx.Document()
    doc.add_heading("GraphRAG Introduction", level=1)
    doc.add_paragraph(
        "Knowledge graphs organize entities and relationships in a structured graph database such as Neo4j."
    )
    doc.add_heading("Vector Search", level=1)
    doc.add_paragraph(
        "Vector search uses high-dimensional embeddings to find semantically similar text passages."
    )
    doc.save(docx_path)
    return docx_path


@pytest.fixture
def sample_pdf_file(tmp_path):
    pdf_path = os.path.join(tmp_path, "sample_doc.pdf")
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((50, 50), "Page 1: GraphRAG combines vector search with Neo4j knowledge graphs.")
    page2 = doc.new_page()
    page2.insert_text((50, 50), "Page 2: Grounded citations ensure answer traceability back to source pages.")
    doc.save(pdf_path)
    doc.close()
    return pdf_path


def test_txt_parser(sample_txt_file):
    registry = ParserRegistry()
    parsed_doc = registry.parse_file(sample_txt_file, doc_id="doc_txt_1")
    assert parsed_doc.doc_id == "doc_txt_1"
    assert parsed_doc.file_type == "txt"
    assert "GraphRAG is a hybrid retrieval architecture" in parsed_doc.full_text


def test_docx_parser(sample_docx_file):
    registry = ParserRegistry()
    parsed_doc = registry.parse_file(sample_docx_file, doc_id="doc_docx_1")
    assert parsed_doc.doc_id == "doc_docx_1"
    assert parsed_doc.file_type == "docx"
    assert len(parsed_doc.pages) == 2
    assert parsed_doc.pages[0].section_header == "GraphRAG Introduction"


def test_pdf_parser(sample_pdf_file):
    registry = ParserRegistry()
    parsed_doc = registry.parse_file(sample_pdf_file, doc_id="doc_pdf_1")
    assert parsed_doc.doc_id == "doc_pdf_1"
    assert parsed_doc.file_type == "pdf"
    assert len(parsed_doc.pages) == 2
    assert parsed_doc.pages[0].page_number == 1
    assert parsed_doc.pages[1].page_number == 2


def test_chunker(sample_txt_file):
    registry = ParserRegistry()
    parsed_doc = registry.parse_file(sample_txt_file, doc_id="doc_txt_1")
    chunks = chunk_parsed_document(parsed_doc, chunk_size=50, chunk_overlap=10)

    assert len(chunks) > 0
    first_chunk = chunks[0]
    assert first_chunk.doc_id == "doc_txt_1"
    assert first_chunk.chunk_index == 0
    assert first_chunk.page_number == 1
    assert first_chunk.start_char_offset >= 0


def test_document_registry(tmp_path):
    db_file = os.path.join(tmp_path, "test_registry.db")
    reg = DocumentRegistry(db_path=db_file)

    doc_info = reg.register_document(
        doc_id="doc_001",
        filename="test.pdf",
        file_path="/tmp/test.pdf",
        file_type="pdf",
    )
    assert doc_info["status"] == "uploaded"

    updated = reg.update_status("doc_001", status="chunked", chunk_count=12)
    assert updated["status"] == "chunked"
    assert updated["chunk_count"] == 12

    doc_list = reg.list_documents()
    assert len(doc_list) == 1
    assert doc_list[0]["doc_id"] == "doc_001"
